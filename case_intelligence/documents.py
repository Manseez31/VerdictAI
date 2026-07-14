"""Document text extraction for case uploads (Feature 1: PDF / DOCX).

Pure, dependency-light text extraction — no LLM, no network — so it is fast and
unit-testable. PDF uses PyMuPDF (already a dependency, same as the report
exporter); DOCX uses python-docx. Extracted text is fed into the Case Analyzer;
we never execute or trust document contents beyond reading their text.
"""

from __future__ import annotations

import io
import logging
from typing import Tuple

logger = logging.getLogger(__name__)

MAX_CHARS = 20000  # generous cap; the analyzer truncates further as needed


class UnsupportedDocument(ValueError):
    """Raised for file types we cannot extract text from."""


def _clean(text: str) -> str:
    # Normalize whitespace/newlines without destroying paragraph structure.
    lines = [ln.rstrip() for ln in (text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    out, blank = [], 0
    for ln in lines:
        if ln.strip():
            out.append(ln)
            blank = 0
        else:
            blank += 1
            if blank <= 1:
                out.append("")
    return "\n".join(out).strip()[:MAX_CHARS]


def extract_text_from_pdf(data: bytes) -> str:
    import fitz  # PyMuPDF

    parts = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            parts.append(page.get_text())
    return _clean("\n".join(parts))


def extract_text_from_docx(data: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs]
    # Include table cell text (contracts/agreements often use tables).
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _clean("\n".join(parts))


def extract_text_from_txt(data: bytes) -> str:
    return _clean(data.decode("utf-8", errors="replace"))


def extract_document_text(filename: str, content_type: str, data: bytes) -> Tuple[str, str]:
    """Route an uploaded file to the right extractor.

    Returns ``(text, kind)`` where kind is 'pdf' | 'docx' | 'txt'. Raises
    UnsupportedDocument for anything else.
    """
    name = (filename or "").lower()
    ctype = (content_type or "").lower()

    if name.endswith(".pdf") or "pdf" in ctype:
        return extract_text_from_pdf(data), "pdf"
    if name.endswith(".docx") or "officedocument.wordprocessingml" in ctype:
        return extract_text_from_docx(data), "docx"
    if name.endswith(".txt") or ctype.startswith("text/plain"):
        return extract_text_from_txt(data), "txt"
    raise UnsupportedDocument("Only PDF, DOCX, or TXT files are supported.")
